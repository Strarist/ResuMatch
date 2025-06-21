import re
import logging
from typing import Dict, List, Any, Optional
from pathlib import Path
import json
from datetime import datetime

try:
    from pdfminer.high_level import extract_text
    from docx import Document
    import PyPDF2
except ImportError as e:
    # Fallback for missing dependencies
    logger = logging.getLogger(__name__)
    logger.warning(f"Some PDF/DOCX parsing libraries not available: {e}")
    extract_text = None
    Document = None
    PyPDF2 = None

logger = logging.getLogger(__name__)

class ResumeParser:
    def __init__(self) -> None:
        # Enhanced skill patterns with more comprehensive coverage
        self.skill_patterns = [
            # Programming Languages
            r'\b(python|java|javascript|typescript|react|angular|vue|node\.js|fastapi|django|flask|spring|express)\b',
            r'\b(c\+\+|c#|php|ruby|go|rust|swift|kotlin|scala|r|matlab)\b',
            
            # Databases
            r'\b(sql|postgresql|mysql|mongodb|redis|cassandra|elasticsearch|dynamodb|oracle)\b',
            
            # Cloud & DevOps
            r'\b(docker|kubernetes|aws|azure|gcp|terraform|ansible|jenkins|git|ci/cd|agile|scrum)\b',
            
            # Frameworks & Libraries
            r'\b(tensorflow|pytorch|scikit-learn|pandas|numpy|matplotlib|seaborn|plotly)\b',
            r'\b(bootstrap|tailwind|sass|less|webpack|babel|eslint|prettier)\b',
            
            # Tools & Platforms
            r'\b(jira|confluence|slack|teams|zoom|figma|sketch|adobe|photoshop|illustrator)\b',
            r'\b(linux|unix|bash|shell|powershell|vim|emacs|vscode|intellij|eclipse)\b',
            
            # Methodologies
            r'\b(agile|scrum|kanban|waterfall|tdd|bdd|devops|microservices|rest|graphql)\b',
            
            # Soft Skills
            r'\b(leadership|communication|teamwork|problem-solving|analytical|creative|collaboration)\b',
            
            # Certifications
            r'\b(aws certified|azure certified|google cloud|pmp|scrum master|agile certified)\b'
        ]
        
        # Experience patterns
        self.experience_patterns = [
            r'(\d{4})\s*[-–]\s*(\d{4}|present|current).*?(senior|junior|lead|principal|staff)?\s*(developer|engineer|manager|analyst|architect|consultant)',
            r'(senior|junior|lead|principal|staff)?\s*(developer|engineer|manager|analyst|architect|consultant).*?(\d{4})\s*[-–]\s*(\d{4}|present|current)',
            r'(software|full-stack|frontend|backend|data|devops|cloud|security)\s*(developer|engineer|architect)',
            r'(project|product|program|technical)\s*(manager|lead|coordinator)'
        ]
        
        # Education patterns
        self.education_patterns = [
            r'(bachelor|master|phd|doctorate|bachelor\'s|master\'s|doctorate\'s).*?(science|arts|engineering|technology|business|administration)',
            r'(b\.s\.|m\.s\.|ph\.d\.|mba|b\.a\.|m\.a\.).*?(computer science|engineering|mathematics|physics|chemistry|biology|economics|finance)',
            r'(computer science|software engineering|information technology|data science|artificial intelligence|machine learning)'
        ]
        
    def parse_resume(self, file_path: str) -> Dict[str, Any]:
        """Parse resume file and extract structured data"""
        try:
            logger.info(f"Starting to parse resume: {file_path}")
            
            # Extract text based on file type
            if file_path.lower().endswith('.pdf'):
                text = self._extract_pdf_text(file_path)
            elif file_path.lower().endswith('.docx'):
                text = self._extract_docx_text(file_path)
            elif file_path.lower().endswith('.doc'):
                text = self._extract_doc_text(file_path)
            else:
                raise ValueError(f"Unsupported file type: {file_path}")
            
            if not text.strip():
                raise ValueError("No text content extracted from file")
            
            logger.info(f"Extracted {len(text)} characters from resume")
            
            # Parse extracted text
            parsed_data = {
                'content': text,
                'skills': self._extract_skills(text),
                'experience': self._extract_experience(text),
                'education': self._extract_education(text),
                'contact_info': self._extract_contact_info(text),
                'summary': self._extract_summary(text),
                'years_of_experience': self._calculate_years_of_experience(text),
                'parsed_at': datetime.utcnow().isoformat()
            }
            
            logger.info(f"Successfully parsed resume with {len(parsed_data['skills'])} skills found")
            return parsed_data
            
        except Exception as e:
            logger.error(f"Error parsing resume {file_path}: {e}")
            raise
    
    def _extract_pdf_text(self, file_path: str) -> str:
        """Extract text from PDF file with multiple fallback methods"""
        text = ""
        
        # Try pdfminer first (most reliable)
        if extract_text is not None:
            try:
                text = extract_text(file_path)
                if text.strip():
                    return text
            except Exception as e:
                logger.warning(f"pdfminer extraction failed: {e}")
        
        # Fallback to PyPDF2
        if PyPDF2 is not None:
            try:
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    text = ""
                    for page in pdf_reader.pages:
                        text += page.extract_text() + "\n"
                if text.strip():
                    return text
            except Exception as e:
                logger.warning(f"PyPDF2 extraction failed: {e}")
        
        raise ImportError("No PDF parsing library available")
    
    def _extract_docx_text(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        if Document is None:
            raise ImportError("python-docx is not installed")
        
        try:
            doc = Document(file_path)
            text_parts = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)
            
            return " ".join(text_parts)
        except Exception as e:
            logger.error(f"Error extracting DOCX text: {e}")
            return ""
    
    def _extract_doc_text(self, file_path: str) -> str:
        """Extract text from DOC file (basic implementation)"""
        # For .doc files, we'll need additional libraries like python-docx2txt
        # For now, return a placeholder
        logger.warning("DOC file parsing not fully implemented")
        return ""
    
    def _extract_skills(self, text: str) -> List[str]:
        """Extract skills from text with improved matching"""
        skills = set()
        text_lower = text.lower()
        
        # Extract skills using patterns
        for pattern in self.skill_patterns:
            matches = re.findall(pattern, text_lower)
            for match in matches:
                if isinstance(match, tuple):
                    skills.update([m for m in match if m])
                else:
                    skills.add(match)
        
        # Additional skill extraction from common sections
        skill_sections = re.findall(r'skills?[:\s]*(.*?)(?=\n\n|\n[A-Z]|experience|education|$)', text, re.IGNORECASE | re.DOTALL)
        for section in skill_sections:
            # Extract individual skills from skill sections
            skill_words = re.findall(r'\b[a-zA-Z][a-zA-Z0-9+#]*\b', section)
            skills.update([word.lower() for word in skill_words if len(word) > 2])
        
        return sorted(list(skills))
    
    def _extract_experience(self, text: str) -> List[Dict[str, str]]:
        """Extract work experience from text with improved patterns"""
        experience_entries = []
        
        for pattern in self.experience_patterns:
            matches = re.finditer(pattern, text, re.IGNORECASE)
            for match in matches:
                entry = {
                    'title': match.group(0).strip(),
                    'duration': '',
                    'company': '',
                    'description': ''
                }
                
                # Try to extract duration
                duration_match = re.search(r'(\d{4})\s*[-–]\s*(\d{4}|present|current)', match.group(0))
                if duration_match:
                    entry['duration'] = f"{duration_match.group(1)} - {duration_match.group(2)}"
                
                experience_entries.append(entry)
        
        return experience_entries
    
    def _extract_education(self, text: str) -> List[str]:
        """Extract education information from text"""
        education = []
        
        for pattern in self.education_patterns:
            matches = re.finditer(pattern, text, re.IGNORECASE)
            for match in matches:
                education.append(match.group(0).strip())
        
        return list(set(education))  # Remove duplicates
    
    def _extract_contact_info(self, text: str) -> Dict[str, str]:
        """Extract contact information from text"""
        contact_info = {}
        
        # Extract email
        email_match = re.search(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', text)
        if email_match:
            contact_info['email'] = email_match.group(0)
        
        # Extract phone (various formats)
        phone_patterns = [
            r'(\+\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}',
            r'(\d{3})[-.\s]?(\d{3})[-.\s]?(\d{4})',
            r'\+?1?[-.\s]?\(?(\d{3})\)?[-.\s]?(\d{3})[-.\s]?(\d{4})'
        ]
        
        for pattern in phone_patterns:
            phone_match = re.search(pattern, text)
            if phone_match:
                contact_info['phone'] = phone_match.group(0)
                break
        
        # Extract LinkedIn
        linkedin_match = re.search(r'linkedin\.com/in/[a-zA-Z0-9-]+', text)
        if linkedin_match:
            contact_info['linkedin'] = linkedin_match.group(0)
        
        return contact_info
    
    def _extract_summary(self, text: str) -> str:
        """Extract summary/objective from text"""
        summary_patterns = [
            r'summary[:\s]*(.*?)(?=\n\n|\n[A-Z]|experience|education|skills|$)',
            r'objective[:\s]*(.*?)(?=\n\n|\n[A-Z]|experience|education|skills|$)',
            r'profile[:\s]*(.*?)(?=\n\n|\n[A-Z]|experience|education|skills|$)',
            r'about[:\s]*(.*?)(?=\n\n|\n[A-Z]|experience|education|skills|$)'
        ]
        
        for pattern in summary_patterns:
            match = re.search(pattern, text, re.IGNORECASE | re.DOTALL)
            if match:
                summary = match.group(1).strip()
                if len(summary) > 20:  # Ensure it's substantial
                    return summary
        
        # Fallback: return first few meaningful sentences
        sentences = re.split(r'[.!?]+', text)
        meaningful_sentences = [s.strip() for s in sentences if len(s.strip()) > 20]
        return meaningful_sentences[0] if meaningful_sentences else ""
    
    def _calculate_years_of_experience(self, text: str) -> int:
        """Calculate approximate years of experience from text"""
        try:
            # Look for year patterns
            year_patterns = [
                r'(\d{4})\s*[-–]\s*(\d{4}|present|current)',
                r'(\d+)\s*(years?|yrs?)\s*of\s*experience',
                r'experience[:\s]*(\d+)\s*(years?|yrs?)'
            ]
            
            total_years = 0
            for pattern in year_patterns:
                matches = re.finditer(pattern, text, re.IGNORECASE)
                for match in matches:
                    if 'years' in match.group(0).lower() or 'yrs' in match.group(0).lower():
                        # Extract number of years
                        year_match = re.search(r'(\d+)', match.group(0))
                        if year_match:
                            total_years = max(total_years, int(year_match.group(1)))
                    else:
                        # Calculate from date range
                        groups = match.groups()
                        if len(groups) >= 2:
                            try:
                                start_year = int(groups[0])
                                end_year = datetime.now().year if groups[1].lower() in ['present', 'current'] else int(groups[1])
                                years = end_year - start_year
                                total_years = max(total_years, years)
                            except ValueError:
                                continue
            
            return total_years
        except Exception as e:
            logger.warning(f"Error calculating years of experience: {e}")
            return 0 